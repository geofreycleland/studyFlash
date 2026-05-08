from docx import Document
import json, re, os, zipfile
src='/mnt/data/questions.docx'
doc=Document(src)
questions=[p.text.strip().replace('‚Äôs',"'s").replace('‚Äú','“').replace('‚Äù','”') for p in doc.paragraphs if p.text.strip()]
# remove any non question? should be 65
opts=[]
for i in range(1,66):
    t=doc.tables[i]
    arr=[]
    for r in t.rows[:4]:
        txt=r.cells[1].text.strip().replace('‚Äôs',"'s").replace('‚Äô',"'").replace('‚Äú','“').replace('‚Äù','”')
        arr.append(txt)
    opts.append(arr)
print(len(questions), len(opts), questions[0], opts[0])
correct = [0,0,0,3,0,1,1,3,1,2,3,1,0,0,1,2,2,3,2,1,3,1,1,[0,1],0,1,[1,3],2,0,2,0,3,3,3,[0,3],[1,2],1,2,0,3,0,[0,2],0,2,1,0,0,[2,4],0,3,0,3,3,0,2,0,0,1,[2,3],0,[1,3],2,1,2,2]
correct_notes = {
1:'RDS automated backups already run daily. Changing the retention period to 30 days meets the requirement with almost no extra work.',
2:'Amazon EMR is built for distributed, parallel processing of large datasets in S3 and can use Redshift data during processing.',
3:'Provisioned IOPS SSD EBS is designed for durable block storage with predictable low latency and consistent I/O performance.',
4:'Predictive scaling uses forecasts and can pre-launch capacity before the weekly batch job, which avoids manual tuning and trend analysis.',
5:'RDS for MySQL keeps MySQL compatibility, and Provisioned IOPS directly addresses heavy write performance needs.',
6:'S3 Standard is good for frequent access in the first 30 days. S3 Standard-IA keeps data durable and available while lowering cost after access drops.',
7:'An SCP in AWS Organizations centrally blocks large EC2 launches across accounts with the least account-by-account work.',
8:'An S3 gateway VPC endpoint keeps S3 traffic on the AWS private network instead of sending it over the internet.',
9:'Kinesis handles ordered streaming spikes, Lambda processes without server management, and DynamoDB provides highly available scalable storage.',
10:'RDS Multi-AZ DB cluster gives high availability and readable instances for reporting without managing Oracle servers yourself.',
11:'A Direct Connect gateway connects Direct Connect to VPCs across Regions at scale without building many separate links.',
12:'CloudFormation automates infrastructure as code, and AWS Config records and audits resource configuration changes.',
13:'Aurora keeps MySQL compatibility and can scale read capacity automatically with Aurora Auto Scaling.',
14:'EFS gives all EC2 instances shared file storage with low lag, which fits content that users constantly update.',
15:'A presigned S3 URL gives temporary access to one object without creating users or opening the whole bucket.',
16:'An Aurora read replica offloads read traffic from the writer, improving write performance during peak read-heavy usage.',
17:'Cross-account IAM roles are the standard secure way to delegate access to a vendor-owned AWS account.',
18:'Secrets Manager integrates with RDS and can rotate database credentials automatically on a schedule.',
19:'Cost and Usage Reports delivered to S3 plus Athena is low-cost and scalable for monthly SQL analysis.',
20:'DAX is a managed, low-latency cache built specifically for DynamoDB and needs only endpoint-level code changes.',
21:'RDS Custom for Oracle provides managed database features while still allowing access to the underlying OS, and a cross-Region read replica supports DR.',
22:'DAX is the lowest-overhead cache for DynamoDB read latency because it is purpose-built for DynamoDB.',
23:'An RDS Multi-AZ DB cluster provides high availability and readable standby instances for near real-time read-only analytics.',
24:'API Gateway and Kinesis Firehose scale ingestion to S3, and Glue can process large raw datasets in S3 with low operational overhead.',
25:'KMS handles secure encryption and decryption, and S3 provides highly available durable encrypted storage.',
26:'Both the SAP application and database show high memory use, so memory optimized EC2 instances fit both tiers.',
27:'Storage Gateway volume gateway supports block storage with local cache, and file gateway supports NFS-style access backed by AWS storage.',
28:'Kinesis Data Firehose can ingest streaming data with low operations and load it into Redshift for SQL analytics.',
29:'EFS is shared file storage that can be mounted by Linux EC2 instances across multiple Availability Zones and grows elastically.',
30:'GuardDuty detects suspicious activity, and EventBridge plus Lambda can automate updates to WAF rules.',
31:'AWS WAF directly protects an ALB from common web attacks like SQL injection and cross-site scripting with low server management.',
32:'A public Direct Connect VIF allows private connectivity from on premises to public AWS services like S3 without using the public internet.',
33:'AWS Control Tower provides account drift notifications for changes such as OU hierarchy drift with low operational overhead.',
34:'CloudFront requires ACM public certificates in us-east-1 for alternate domain names, and public ACM certificates are free.',
35:'SQS decouples purchase requests from processing, Auto Scaling workers handle spikes, and RDS Proxy protects Aurora from connection storms.',
36:'Textract extracts text from scanned documents, Comprehend Medical extracts medical entities, and Athena can query data stored in S3.',
37:'Multi-AZ Auto Scaling plus Multi-AZ database deployment and RDS Proxy improves availability with low operational effort.',
38:'A scheduled Lambda can check DynamoDB daily, and SNS can send email notifications without managing servers.',
39:'ECS on Fargate runs existing containers with little code change, Auto Scaling, and no EC2 server management.',
40:'Pre-created DR Auto Scaling and load balancing plus DynamoDB global tables gives the lowest downtime; automated DNS switching routes users to DR.',
41:'SageMaker is the managed ML service for building and training models, and QuickSight is the managed BI/dashboard service.',
42:'DynamoDB uses a gateway VPC endpoint, and route table entries send DynamoDB traffic through that endpoint privately.',
43:'QuickSight dashboards can visualize Cost and Usage Report data queried from S3 by Athena.',
44:'SQS durably stores jobs and decouples producers from workers. Auto Scaling can scale workers based on queue depth.',
45:'Session Manager provides secure shell-like access without shared SSH keys or bastion hosts.',
46:'AWS Backup centrally manages DynamoDB backup schedules and long retention policies such as 7 years.',
47:'An active-passive DR setup with Route 53 and an Aurora replica in another Region meets 30-minute RTO/RPO better than restore-only designs.',
48:'Firehose provides serverless streaming ingestion, and ECS Fargate handles long-running 30-minute processing without Lambda time limits.',
49:'ACM with DNS validation automates public TLS certificate issuance and renewal for CloudFront.',
50:'Textract extracts text from PDFs/JPEGs, and Comprehend Medical identifies PHI with less custom code.',
51:'AMI copies plus CloudFormation templates are a pilot-light style DR approach with few resources running normally and faster recovery than manual scripts.',
52:'SQS uses an interface VPC endpoint. Placing it in private subnets with security groups allows private EC2 access.',
53:'An S3 gateway endpoint lets EC2 instances call S3 APIs privately without internet routing.',
54:'Amazon Macie is the managed service for discovering sensitive data such as PII in S3 and must be enabled per Region.',
55:'A scheduled ECS task on Fargate can run for an hour with known CPU/memory and avoids managing EC2 servers.',
56:'Cognito handles authentication, Lambda@Edge supports low-latency authorization near users, and CloudFront serves globally.',
57:'A Compute Savings Plan discounts compute while allowing changes across instance families, sizes, OS, and Regions.',
58:'Transit Gateway connects multiple VPCs and Direct Connect efficiently, avoiding separate connections per VPC.',
59:'To target only nonproduction, attach the SCP directly to the three nonproduction accounts or put them in a new OU and attach the SCP there.',
60:'Active DR infrastructure plus DynamoDB global tables and Route 53 failover gives the least downtime.',
61:'For high-use RDS On-Demand databases, Trusted Advisor RDS Reserved Instance Optimization in the management account helps find cost savings across the organization.',
62:'AWS Config has a managed rule for ACM certificate expiration, and EventBridge/SNS can notify when resources become noncompliant.',
63:'Secrets Manager is designed for database password storage and automatic rotation with the least operational overhead.',
64:'AWS Glue is the managed ETL service for converting CSV data and loading processed data into Redshift with low operations.',
65:'CloudFront signed URLs are meant to grant temporary, controlled access to specific premium content such as rentals or downloads.'
}
service_defs = {
'Amazon RDS':'Managed relational database service for engines like MySQL, PostgreSQL, Oracle, and SQL Server.', 'RDS':'Managed relational database service for relational databases.',
'Amazon EMR':'Managed big data platform for distributed processing with Spark, Hadoop, and related tools.', 'Amazon S3':'Object storage for durable, scalable storage of files and data.',
'Amazon Redshift':'Managed data warehouse for SQL analytics at scale.', 'Amazon EBS':'Block storage volumes for EC2 instances.', 'ElastiCache':'Managed in-memory caching service for Redis or Memcached.',
'EC2':'Virtual servers in AWS.', 'Auto Scaling':'Automatically adjusts capacity to match demand.', 'EventBridge':'Event bus and scheduler for routing events and running scheduled tasks.', 'AWS Lambda':'Serverless compute that runs code without managing servers.',
'CloudWatch':'Monitoring and metrics service for AWS resources.', 'Amazon EFS':'Managed shared NFS file storage for Linux workloads.', 'DynamoDB':'Managed NoSQL key-value and document database.', 'DAX':'DynamoDB Accelerator, an in-memory cache for DynamoDB.',
'AWS Organizations':'Central management for multiple AWS accounts.', 'SCP':'Service control policy that sets permission guardrails in AWS Organizations.', 'VPC endpoint':'Private connection from a VPC to supported AWS services.', 'Kinesis Data Streams':'Managed real-time streaming service for ordered event data.', 'Kinesis Data Firehose':'Managed delivery stream that loads streaming data into destinations like S3 and Redshift.', 'SNS':'Pub/sub notification service that can fan out messages and send emails.', 'SQS':'Managed message queue for decoupling systems.', 'DMS':'Database Migration Service for moving databases to AWS.', 'CloudFormation':'Infrastructure as code service for provisioning AWS resources.', 'AWS Config':'Tracks resource configuration and compliance over time.', 'Aurora':'AWS cloud-native relational database compatible with MySQL and PostgreSQL.', 'Secrets Manager':'Secure secret storage with automatic rotation support.', 'Athena':'Serverless SQL query service for data in S3.', 'RDS Proxy':'Managed database proxy that pools and manages application database connections.', 'Textract':'Extracts text and structured data from scanned documents and images.', 'Comprehend Medical':'Detects medical information and PHI from text.', 'GuardDuty':'Threat detection service for AWS accounts and workloads.', 'AWS WAF':'Web application firewall for filtering HTTP/S traffic.', 'Direct Connect':'Dedicated private network connection from on premises to AWS.', 'Control Tower':'Governance service for setting up and managing multi-account AWS environments.', 'ACM':'AWS Certificate Manager for provisioning and renewing TLS certificates.', 'QuickSight':'Managed business intelligence and dashboard service.', 'SageMaker':'Managed machine learning service for building, training, and deploying models.', 'Systems Manager Session Manager':'Secure browser/CLI access to instances without SSH keys.', 'AWS Backup':'Centralized backup service for AWS resources.', 'Route 53':'DNS and health-check service.', 'Fargate':'Serverless compute engine for containers on ECS or EKS.', 'ECS':'Container orchestration service for running Docker containers.', 'CloudFront':'Global content delivery network.', 'Cognito':'Managed user sign-up, sign-in, and authentication.', 'Lambda@Edge':'Runs Lambda code at CloudFront edge locations.', 'Macie':'Sensitive data discovery service for S3.', 'Storage Gateway':'Hybrid cloud storage service connecting on-premises apps to AWS storage.', 'Glue':'Managed ETL and data integration service.', 'KMS':'Key Management Service for creating and controlling encryption keys.', 'Trusted Advisor':'AWS best-practice checks for cost, security, performance, and more.'
}
def letter(idx): return chr(65+idx)
def services_for(text):
    found=[]
    alltext=text
    for k,v in service_defs.items():
        if re.search(r'\b'+re.escape(k)+r'\b', alltext, flags=re.I):
            found.append({'name':k,'definition':v})
    # de-dupe by definition/name
    seen=set(); out=[]
    for f in found:
        key=f['name'].lower()
        if key not in seen:
            seen.add(key); out.append(f)
    return out[:10]
def wrong_reason(option, qnum):
    s=option
    low=s.lower()
    if 'manual' in low or 'cli' in low or 'custom' in low or 'script' in low or 'cron' in low:
        return 'This adds more manual work or custom maintenance, so it is not the least operational overhead choice.'
    if 'ec2' in low and ('fleet' in low or 'instances' in low):
        return 'This can work in some designs, but it requires more server management than the managed/serverless option.'
    if 'single-az' in low:
        return 'Single-AZ does not meet the high availability requirement.'
    if 'internet' in low or 'nat gateway' in low:
        return 'This does not keep traffic fully private to the AWS service or adds unnecessary network cost.'
    if 'redshift' in low and qnum not in [19,28,43,64]:
        return 'Redshift is mainly for analytics/data warehousing, not the best fit for the application requirement here.'
    if 's3 glacier' in low or 'glacier' in low:
        return 'Glacier classes are for archive access and are not ideal when objects must be immediately available.'
    if 'one zone' in low:
        return 'One Zone storage is cheaper but less resilient because it stores data in one Availability Zone.'
    if 'parameter store' in low:
        return 'Parameter Store can store values, but it is not the simplest fully managed rotation choice for this database password requirement.'
    if 'sns' in low and qnum in [44]:
        return 'SNS is pub/sub and does not durably hold work items like a queue.'
    if 'sqs' in low and qnum in [38]:
        return 'SQS queues messages but does not directly send email notifications.'
    return 'This choice misses a key requirement, is less managed, or is less cost-effective than the correct answer.'
items=[]
for i,(q,o,c) in enumerate(zip(questions,opts,correct), start=1):
    corr = c if isinstance(c,list) else [c]
    alltext=q+' '+' '.join(o)
    why_wrong=[]
    for idx,opt in enumerate(o):
        if idx not in corr:
            why_wrong.append({'choice':letter(idx),'text':opt,'reason':wrong_reason(opt,i)})
    items.append({'id':i,'question':q,'options':o,'correct':corr,'answerText':', '.join(letter(x) for x in corr),'whyRight':correct_notes[i],'whyWrong':why_wrong,'services':services_for(alltext)})
with open('/mnt/data/aws-study-site/questions.js','w') as f:
    f.write('const QUESTIONS = ' + json.dumps(items, indent=2) + ';')
